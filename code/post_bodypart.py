"""
The script for post processing body part regression
1) Implement lenear regression and robust regression for all the scores
2) Generate a report for each volume in its directory
3) Crop the data if needed
Yucheng Tang
OCT 2018

"""
import os
import numpy as np
import nibabel as nb
from nibabel.processing import resample_from_to 
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image

from docx import Document
from docx.shared import Inches
from sklearn import linear_model, datasets
import matplotlib.pyplot as plt
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from PIL import Image, ImageDraw

class post_bodypart_regression(object):
    def __init__(self, args):
        """Initialize and post processing ."""
        self.data_dir = args.data_dir
        self.cropped_dir = args.cropped_dir
        self.txt_info_dir = args.txt_info_dir


        self.score_file = os.path.join(self.txt_info_dir, 'BPR_result_list.txt')

        self.score_interval = []
        self.score_interval.append(float(args.score_interval.split(' ')[0]))
        self.score_interval.append(float(args.score_interval.split(' ')[1]))
        
    def get_slice(self,datascore_file):
        with open(datascore_file) as f:
            content = f.readlines()
        content = [x.strip() for x in content]
        name2idx_scores = {}
        name2slice = {}
        name2mb = {}
        for item in content:
            image_name = item.split('_idx')[0]
            if image_name not in name2idx_scores:
                name2idx_scores[image_name] = [[],[]]
            slice_idx = int(item.split('.png')[0].split('_idx')[1])
            idx_score = float(item.split('\t')[1])
            name2idx_scores[image_name][0].append(slice_idx)
            name2idx_scores[image_name][1].append(idx_score)   
        for item in name2idx_scores:
            name2mb[item] = [] 
            x_vals = name2idx_scores[item][0]
            y_vals = name2idx_scores[item][1]
        return name2idx_scores

    def slope_intercept(self, x_val, y_val):
        """
        Linear regression 
        input: vector x, vector y, size must be the same
        return: slope and intercept
        """
        x_val = x_val.reshape(x_val.shape[0])
        x = np.array(x_val)
        y = np.array(y_val)
        m = (((np.mean(x) * np.mean(y))) - np.mean(x*y)) / ((np.mean(x) * np.mean(x) -  np.mean(x*x)))
        # m = round(m, 2)
        b = (np.mean(y) - np.mean(x) * m)
        # b = round(b,2)
        return m,b
    
    def processing(self):
        count = 0
        crop_index_txt_file = os.path.join(self.txt_info_dir, 'bpr_crop_index.txt')
        index_file = open(crop_index_txt_file, 'w')
        image_dir = os.path.join(self.data_dir, 'images')
        soft_image_dir = os.path.join(self.data_dir, 'soft_images')
        label_dir = os.path.join(self.data_dir, 'labels')
        BPR_reports_dir = os.path.join(self.data_dir, 'BPR_reports')
        os.system('rm -rf \"{}\"'.format(BPR_reports_dir))
        os.makedirs(BPR_reports_dir)
        # temp_dir = '/nfs/masi/tangy5/bpr_pipeline/INPUTS/temp'
        for volume in os.listdir(soft_image_dir):
            if volume.endswith('.nii.gz'):
                count += 1
                image_path = os.path.join(soft_image_dir, volume)
                imgnb = nb.load(image_path)
                imgnp = np.array(imgnb.dataobj)
                print('processing {}'.format(volume))

                document  = Document()
                # Add page settings for margin
                sections = document.sections
                for section in sections:
                    section.top_margin = Cm(2)
                    section.bottom_margin = Cm(2)
                    section.left_margin = Cm(2)
                    section.right_margin = Cm(2)
                # Add heading and font style
                document.add_heading('{}'.format(volume), 0)    
                p = document.add_paragraph().add_run()
                p.add_text('Body part regression report')
                p.bold = True
                p.italic = True
                font_p = p.font
                font_p.name = 'Calibri'
                font_p.size = Pt(20)
                # Add Volume information table
                document.add_heading('Volume Info', level=1)
                table = document.add_table(rows=1, cols=3)
                table.style = 'Light List Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Dimension'
                hdr_cells[1].text = 'Pixdim'
                hdr_cells[2].text = 'Data Type'
                dim = '[{} {} {}]'.format(imgnb.header['dim'][1],imgnb.header['dim'][2],imgnb.header['dim'][3])
                pixdim = '[{} {} {}]'.format(imgnb.header['pixdim'][1],imgnb.header['pixdim'][2],imgnb.header['pixdim'][3])
                dtype = '{}'.format(imgnb.header['datatype'].dtype)
                row_cells = table.add_row().cells
                row_cells[0].text = dim
                row_cells[1].text = pixdim
                row_cells[2].text = dtype

                document.add_paragraph('', style='Intense Quote')


                # Add scatter and regression figure
                document.add_paragraph(
                'Regression and Scatter result', style='List Bullet'
                )

                # Using robust fit and linear for regression, create a scatter-regression figure
                name2idx_scores = self.get_slice(self.score_file)
                # print(name2idx_scores)

                try: 
                    x_vals = name2idx_scores[volume][0]
                    y_vals = name2idx_scores[volume][1]
                except KeyError:
                    print('Error Image, skipping')
                    continue
                # Make x_vals [n,1] size. y_vals [n,] to numnp
                x_vals = np.array(x_vals)
                x_vals = x_vals.reshape((len(x_vals), 1))
                y_vals = np.array(y_vals)
                lr = linear_model.LinearRegression()
                ransac = linear_model.RANSACRegressor()
                lr.fit(x_vals, y_vals)

                try:
                    ransac.fit(x_vals, y_vals)
                    line_x = np.arange(x_vals.min(), x_vals.max())[:, np.newaxis]
                    inlier_mask = ransac.inlier_mask_
                    outlier_mask = np.logical_not(inlier_mask)
                    line_y_ransac = ransac.predict(line_x)
                    line_y = lr.predict(line_x)
                except ValueError:
                    print('Error Image, skipping')
                    continue
                lw = 2
                fig = plt.figure()
                # ax = plt.gca()
                plt.scatter(x_vals[inlier_mask], y_vals[inlier_mask], color='yellowgreen', marker='.',
                            label='Inliers')
                plt.scatter(x_vals[outlier_mask], y_vals[outlier_mask], color='gold', marker='.',
                            label='Outliers')
                plt.plot(line_x, line_y, color='navy', linewidth=lw, label='Linear regressor')
                plt.plot(line_x, line_y_ransac, color='cornflowerblue', linewidth=lw,
                        label='RANSAC regressor')
                plt.legend(loc='upper right')
                plt.xlabel('Slices')
                plt.ylabel('Scores')

                sca_reg_fig = os.path.join(BPR_reports_dir, '{}_sca_reg.png'.format(volume))
                fig.savefig(sca_reg_fig)
                try:
                    m1, b1 = self.slope_intercept(line_x, line_y)
                    m2, b2 = self.slope_intercept(line_x, line_y_ransac)
                    z1_max = int(round((float(self.score_interval[0]) - b1) / m1))
                    z1_min = int(round((float(self.score_interval[1]) - b1) / m1))
                    z2_max = int(round((float(self.score_interval[0]) - b2) / m2))
                    z2_min = int(round((float(self.score_interval[1]) - b2) / m2))
                except ValueError:
                    print('Error Image, skipping')
                    continue

                if z1_min < 0:
                    z1_min = 0
                if z2_min < 0:
                    z2_min = 0
                if z1_max > imgnb.header['dim'][3]:
                    z1_max = imgnb.header['dim'][3]
                if z2_max > imgnb.header['dim'][3]:
                    z2_max = imgnb.header['dim'][3]
                if z1_max <=0 or z2_max <= 0:
                    print('QA failed, skipping {}'.format(volume))
                    continue
                # Add axial image
                try:
                    img_axial_min = Image.fromarray(imgnp[:,:,z2_min] * 255.0).convert('L').rotate(90)
                    img_axial_max = Image.fromarray(imgnp[:,:,z2_max - 1] * 255.0).convert('L').rotate(90)
                except IndexError:
                    print('QA Failed, Skipping {}'.format(volume))
                    continue
                image_axial_min = os.path.join(BPR_reports_dir,  '{}_axial_min.png'.format(volume))
                image_axial_max = os.path.join(BPR_reports_dir,  '{}_axial_max.png'.format(volume))
                img_axial_min.save(image_axial_min)   
                img_axial_max.save(image_axial_max)   

                pri = document.add_paragraph()
                run3 = pri.add_run()
                run3.add_picture(sca_reg_fig,width=Inches(3))
                run3.add_text('Min')
                run3.add_picture(image_axial_min, width=Inches(1))
                run3.add_text('Max')
                run3.add_picture(image_axial_max, width=Inches(1))
                
                LR_reg_result = document.add_paragraph('Linear Regression :The required whole abdominal axial dimension score from {} to {} for {} : Start slice {}, End slice {}'.format(self.score_interval[0],\
                    self.score_interval[1], volume, z1_min, z1_max), style='List Number').add_run()
                font_lr_reg = LR_reg_result.font
                font_lr_reg.name = 'Calibri'
                font_lr_reg.size = Pt(12)

                RR_reg_result = document.add_paragraph('Robust Linear Regression :The required whole abodminal axial dimension score from {} to {} for {} : Start slice {}, End slice {}'.format(self.score_interval[0],\
                    self.score_interval[1], volume, z2_min, z2_max),style='List Number').add_run()
                font_rr_reg = RR_reg_result.font
                font_rr_reg.name = 'Calibri'
                font_rr_reg.size = Pt(12)

                # Implement crop here
                x_index = imgnp.shape[0]
                y_index = imgnp.shape[1]
                z_index = imgnp.shape[2]


                index_file.write(volume + ' ' + str(z2_min) + ' ' + str(z2_max) + ' ' + str(imgnb.header['dim'][3]) + '\n')

                #Find corresponding unnormlized volume and crop, save both volume for each case
                soft_output_nii_path = os.path.join(self.cropped_dir, 'soft_images', volume)
                out_nii_name = volume.replace('.nii.gz','.nii.gz')
                output_nii_path = os.path.join(self.cropped_dir, 'images', out_nii_name)

                if os.path.isfile(soft_output_nii_path) and os.path.isfile(output_nii_path):
                    print('[{}] Skipping {}'.format(count, volume))
                    continue

                print('1-Cropping soft tissue windowed volume {} from z index [{},{}] to [{},{}]'.format(volume, 0, z_index,\
                        z2_min, z2_max ))
                


                # original_image_path = os.path.join(temp_dir, out_nii_name)
                image_volume = volume.replace('_soft.nii.gz', '.nii.gz')
                original_image_path = os.path.join(image_dir, image_volume)


                os.system('fslroi \"{}\" \"{}\" {} {} {} {} {} {}'.format(image_path,
                        soft_output_nii_path, 0, x_index, 0, y_index, int(z2_min), int(z2_max-z2_min)))
                
                print('2-Original processing {}'.format(original_image_path))
                os.system('fslroi \"{}\" \"{}\" {} {} {} {} {} {}'.format(original_image_path,
                        output_nii_path, 0, x_index, 0, y_index, int(z2_min), int(z2_max-z2_min)))
                                                      
                # Crop the segmenation nii
                # label_path = os.path.join(label_dir, volume)
                # output_label_dir = os.path.join(self.cropped_dir, 'labels')
                # if not os.path.isdir(output_label_dir):
                #     os.makedirs(output_label_dir)
                # output_label_path = os.path.join(output_label_dir, volume)

                # os.system('fslroi \"{}\" \"{}\" {} {} {} {} {} {}'.format(label_path,
                #         output_label_path, 0, x_index, 0, y_index, int(z2_min), int(z2_max-z2_min+1)))
                # print('3-Label processing {}'.format(volume))


                print('Cropping session complete!')
                
                # New section
                document.add_paragraph(
                '2D views in Soft tissue window within HU [-175, 275]', style='List Bullet'
                )
                # Save 2d images

                img_sagittal = Image.fromarray(imgnp[256,:,:] * 255.0).convert('L').rotate(180)
                img_sagittal_np = np.array(img_sagittal)
                img_sagittal_np = np.transpose(img_sagittal_np, (1, 0))
                img_sagittal = Image.fromarray(img_sagittal_np)

                line1 = ImageDraw.Draw(img_sagittal)
                line1.line((0, img_sagittal.size[1] - z2_min, 511, img_sagittal.size[1] - z2_min), fill=255, width=3)
                line2 = ImageDraw.Draw(img_sagittal)
                line2.line((0, img_sagittal.size[1] - z2_max, 511, img_sagittal.size[1] - z2_max), fill=255, width=3)

                img_coronal = Image.fromarray(imgnp[:,256,:] * 255.0).convert('L').rotate(180)
                img_coronal_np = np.array(img_coronal)
                img_coronal_np = np.transpose(img_coronal_np, (1, 0))
                img_coronal = Image.fromarray(img_coronal_np)

                line3 = ImageDraw.Draw(img_coronal)
                line3.line((0, img_coronal.size[1] - z2_min, 511, img_coronal.size[1] - z2_min), fill=255, width=3)
                line4 = ImageDraw.Draw(img_coronal)
                line4.line((0, img_coronal.size[1] - z2_max, 511, img_coronal.size[1] - z2_max), fill=255, width=3)



                image_coronal = os.path.join(BPR_reports_dir,  '{}_coronal.png'.format(volume))
                image_sagittal = os.path.join(BPR_reports_dir, '{}_sagittal.png'.format(volume))
                img_coronal.save(image_coronal)
                img_sagittal.save(image_sagittal)

                pi = document.add_paragraph()
                run = pi.add_run()
                run.add_text('Coronal')
                run.add_picture(image_coronal, width=Inches(1.25), height=Inches(1.25))
                run.add_text('Sagittal')
                run.add_picture(image_sagittal, width=Inches(1.25),height=Inches(1.25))        


                document.add_page_break()
                doc_name = '{}_report.docx'.format(volume)
                doc_file = os.path.join(BPR_reports_dir, doc_name)
                print('[{}] Saving {} to {}'.format(count, doc_name, doc_file))
                document.save(doc_file)
        index_file.close()
